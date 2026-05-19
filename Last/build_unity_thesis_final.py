from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "论文-基于Unity的类银河恶魔城游戏设计与实现-按Last要求重制.docx"

TITLE = "基于Unity的类银河恶魔城游戏设计与实现"
STUDENT = "吴奇迅"
COLLEGE = "计算机科学学院"
TEACHER = "陈应霞"
MAJOR = "软件工程"

ABSTRACT_CN = [
    "随着独立游戏开发工具链的成熟，二维动作探索类游戏逐渐成为高校游戏开发毕业设计中的重要选题方向。类银河恶魔城游戏兼具动作操作、地图探索、能力解锁与成长反馈等特点，既能够体现完整的系统设计能力，也能较好地展示程序架构组织与玩法实现过程。基于此，本文围绕一个使用 Unity 引擎开发的二维类银河恶魔城游戏项目展开设计与实现研究。",
    "本项目采用 Unity 2022.3 LTS 作为开发平台，使用 C# 作为核心脚本语言，围绕角色控制、战斗系统、敌人 AI、地图探索、能力解锁、背包装备、属性成长与多档位存档等模块展开实现。系统在角色层面构建了有限状态机驱动的行为框架，完成待机、移动、跳跃、空中、冲刺、贴墙滑行、墙跳、近战攻击、远程攻击、受击、回声施法和回声换位等状态切换；在关卡层面基于 Tilemap 设计主场景与 Boss 场景，并通过能力门控、区域发现和特殊通道机制增强探索性；在成长层面通过经验、等级、金币、装备和属性修正构建角色成长闭环。",
    "在工程实现上，项目按照玩家模块、敌人模块、成长模块、界面模块、音频模块和存档模块进行组织，通过接口、事件与组件引用实现模块协同。存档系统采用 JSON 序列化方案，可保存玩家位置、生命值、经验等级、已解锁能力、背包装备、已发现区域与已移除场景对象等信息，支持三档位保存、读取和删除操作。",
    "测试结果表明，该系统能够完成从主菜单进入游戏、角色探索、战斗、能力获取、界面管理到存档恢复的完整流程，具备较好的可玩性和系统完整性。本文的研究与实现表明，Unity 与有限状态机结合的模块化方案适合中小型二维动作游戏项目开发，也可为同类课题提供参考。"
]

ABSTRACT_EN = [
    "With the maturity of independent game development tools, 2D action-exploration games have become an important topic in undergraduate graduation projects. Metroidvania-style games combine action control, map exploration, ability unlocking, and progression feedback, making them suitable for demonstrating both gameplay design and software engineering capabilities. Based on this background, this thesis studies the design and implementation of a Unity-based 2D Metroidvania game project.",
    "The project is developed with Unity 2022.3 LTS and C#. Its core modules include character control, combat, enemy AI, map exploration, ability unlocking, inventory and equipment management, character progression, and multi-slot save/load support. A finite-state-machine-driven behavior framework is built for the player, covering idle, move, jump, air, dash, wall slide, wall jump, melee attack, ranged attack, damaged response, echo casting, and echo swapping states. The level structure is created with Tilemap and expanded through ability gates, discoverable areas, and special passages.",
    "From the engineering perspective, the project is organized into player, enemy, progression, UI, audio, and save modules. These modules collaborate through interfaces, events, and component references. The save system uses JSON serialization and supports three save slots, storing player position, health, level data, unlocked abilities, inventory and equipment state, discovered areas, and removed scene objects.",
    "The test results show that the system can complete the full gameplay loop from menu entry to exploration, combat, ability acquisition, UI management, and save restoration. The study demonstrates that a modular architecture based on Unity and finite state machines is suitable for small and medium 2D action game projects and can serve as a reference for similar graduation designs."
]

KEYWORDS_CN = "Unity；类银河恶魔城；有限状态机；Tilemap；角色成长；游戏存档"
KEYWORDS_EN = "Unity; Metroidvania; Finite State Machine; Tilemap; Character Progression; Save System"

REFERENCES = [
    "［1］Nystrom R．Game Programming Patterns［M］．USA：Genever Benning，2014．",
    "［2］Gregory J．Game Engine Architecture［M］．3rd ed．Boca Raton：CRC Press，2018．",
    "［3］Schell J．The Art of Game Design：A Book of Lenses［M］．3rd ed．Boca Raton：CRC Press，2019．",
    "［4］Swink S．Game Feel：A Game Designer's Guide to Virtual Sensation［M］．Boca Raton：CRC Press，2008．",
    "［5］Bycer J．Game Design Deep Dive：Metroidvania［M］．CRC Press，2025．",
    "［6］Unity Technologies．Unity User Manual［EB/OL］．",
    "［7］Unity Technologies．Tilemap System Documentation［EB/OL］．",
    "［8］Unity Technologies．Animation State Machines Documentation［EB/OL］．",
    "［9］Oliveira B P，Franco A O R，da Silva J W F，et al．A Framework for Metroidvania Games［C］．Proceedings of SBGames 2020，2020．",
    "［10］Wahlberg T．Blockades in the Metroidvania Genre of Games：An Examination of Backtracking［D］．2015．",
    "［11］Yousuf M H，Loutfouz Z．NCCollab：Collaborative Behavior Tree Authoring in Game Development［J］．Multimedia Tools and Applications，2022：31-38．",
    "［12］David B．Game Development Patterns with Unity 2021［M］．Packt Publishing，2021．"
]

ACKS = [
    "在本次毕业设计与论文撰写过程中，我得到了指导教师、同学和家人的支持与帮助。在论文完成之际，谨向所有关心和帮助过我的老师与亲友表示诚挚感谢。",
    "首先，衷心感谢陈应霞老师在课题选题、系统设计、论文修改和整体进度安排方面给予的耐心指导。老师严谨认真的治学态度让我在项目实现与论文写作中受益匪浅。",
    "其次，感谢计算机科学学院各位任课教师在大学期间的培养，使我在程序设计、软件工程、数据结构和计算机图形相关课程中打下了较扎实的基础，为本课题的完成提供了知识支持。",
    "同时，感谢同学们在项目调试、玩法体验和论文整理过程中提供的交流与建议，也感谢家人在整个学习阶段给予的理解、鼓励与陪伴。",
    "由于个人水平和时间有限，论文中仍可能存在不足之处，恳请各位老师批评指正。"
]


SECTIONS = [
    ("h1", "1 绪论"),
    ("h2", "1.1 选题背景及意义"),
    ("p", "类银河恶魔城游戏是一类以非线性地图探索、能力驱动推进和动作反馈为核心的二维游戏类型。与传统线性闯关游戏相比，该类作品更强调地图结构设计与角色成长系统之间的联动关系。玩家通过获得新的能力不断拓展可探索区域，在回访旧区域的过程中形成“探索—解锁—再探索”的循环体验。正因为这种玩法兼具程序实现复杂度和系统组织完整度，所以它非常适合作为游戏开发方向的毕业设计课题。"),
    ("p", "结合任务书要求，本课题以 Unity 为开发平台，围绕高精度角色控制系统、战斗与资源循环、敌人 AI 设计、场景探索和多档位存档等内容展开实现。该项目不仅要求完成一个具备可玩性的游戏原型，还需要从软件工程视角说明如何利用模块化方法组织玩家、敌人、成长、界面与数据持久化等多个系统。"),
    ("p", "从实践意义上看，本课题有助于验证 Unity 在二维动作探索类项目中的适用性，梳理有限状态机、Tilemap、事件驱动通信和 JSON 数据持久化在游戏开发中的协同方式；从教学意义上看，该项目能够较完整地体现需求分析、系统设计、编码实现和测试验证的全过程。"),
    ("h2", "1.2 国内外现状分析"),
    ("p", "国外在类银河恶魔城游戏设计方面已有较成熟的发展经验。经典作品与现代独立游戏普遍将能力门控、地图回访、战斗反馈和场景叙事作为核心设计对象。相关研究与实践不仅关注关卡布局，也强调动作手感、成长节奏和敌人行为设计在整体体验中的作用。"),
    ("p", "国内对该方向的研究起步相对较晚，但随着 Unity 的普及和独立游戏开发环境的成熟，越来越多的项目开始关注二维平台动作与探索系统的综合实现。当前多数教学与毕业设计项目通常聚焦于角色控制、关卡编辑、背包存档和敌人 AI 等内容，并逐步从单一功能演示转向更完整的系统闭环。"),
    ("h2", "1.3 本文主要工作"),
    ("p", "本文围绕当前 Unity 项目，重点完成以下工作：一是构建有限状态机驱动的角色控制系统，实现移动、跳跃、冲刺、墙面交互与多种攻击状态；二是设计普通敌人、虫类敌人和迷你 Boss 的状态机行为；三是实现能力解锁、区域发现、回声换位和门控探索机制；四是完成经验、等级、金币、装备和属性成长系统；五是实现程序化 UI 和三档位 JSON 存档功能，并结合项目实际运行情况进行测试分析。"),
    ("h2", "1.4 论文组织结构"),
    ("p", "全文共分为七章。第一章介绍课题背景、研究现状与论文工作内容；第二章介绍系统开发技术；第三章进行系统分析；第四章说明系统总体设计；第五章详细阐述关键模块实现；第六章给出系统测试与结果分析；第七章对全文进行总结并提出后续展望。"),

    ("h1", "2 系统开发技术"),
    ("h2", "2.1 Unity 引擎"),
    ("p", "Unity 是当前二维与三维交互内容开发中应用广泛的引擎之一。项目使用 Unity 2022.3.12f1c1 版本进行开发，其集成了场景编辑、物理模拟、动画控制、资源管理和 UI 构建等功能，能够较好满足毕业设计对开发效率和系统完整性的要求。"),
    ("h2", "2.2 C# 语言与面向对象设计"),
    ("p", "C# 是 Unity 的核心脚本语言，具备清晰的面向对象表达能力。项目中通过类继承、接口、多态和事件机制组织游戏逻辑。例如，玩家与敌人共享基础实体能力，战斗判定通过 IDamageable 接口解耦，仇恨目标通过 IAggroTarget 描述，从而使系统结构更加清晰。"),
    ("h2", "2.3 有限状态机技术"),
    ("p", "有限状态机适合处理动作游戏中大量离散的行为切换问题。本项目在玩家系统和敌人系统中都使用了状态机。每个状态类负责进入、更新和退出逻辑，状态机统一完成状态切换，使复杂条件判断被拆分到多个职责清晰的状态对象中。"),
    ("h2", "2.4 Unity 2D 物理与动画系统"),
    ("p", "项目依赖 Rigidbody2D、Collider2D 和 Physics2D 完成移动、碰撞、受击和攻击范围检测，并结合 Animator 与动画控制器驱动角色和敌人的动作表现。物理系统为横版动作体验提供基础，动画系统则用于强化状态反馈与视觉连贯性。"),
    ("h2", "2.5 Tilemap 与 JSON 序列化"),
    ("p", "在地图构建方面，项目基于 Tilemap 组织主场景与 Boss 场景，通过瓦片绘制完成地形、装饰与未探索区域的布置。在数据持久化方面，系统使用 JsonUtility 将游戏进度序列化为 JSON 文件，以实现多档位存档管理。"),

    ("h1", "3 系统分析"),
    ("h2", "3.1 可行性分析"),
    ("h3", "3.1.1 技术可行性"),
    ("p", "本项目采用的 Unity、C#、Tilemap、Animator 和 JsonUtility 等技术均具备成熟的开发文档与实践案例，适合完成二维动作探索类游戏原型。项目脚本已划分为 Player、Progression、Save、UI、Audio 和 Combat 等模块，说明技术路径具备良好的落地基础。"),
    ("h3", "3.1.2 经济可行性"),
    ("p", "毕业设计项目主要依赖个人开发和学校现有实验环境，不需要额外采购复杂商业中间件。Unity 个人版即可支撑开发流程，项目资源以自制与基础素材整合为主，因此总体开发成本较低。"),
    ("h3", "3.1.3 操作可行性"),
    ("p", "系统运行在 Windows 平台和 Unity 编辑器环境中，操作方式明确，菜单、角色控制和存档流程均较为直观。对于课题展示和答辩演示而言，该项目具有较好的可操作性。"),
    ("h2", "3.2 功能需求分析"),
    ("p", "结合任务书和现有项目实现，系统功能需求可归纳为角色控制、战斗系统、敌人 AI、地图探索、成长系统、界面交互和存档管理七个方面。"),
    ("table_req", ""),
    ("h2", "3.3 非功能需求分析"),
    ("p", "除了基本功能外，系统还需要满足以下非功能需求：第一，代码结构要清晰，便于后续扩展新能力、新敌人和新关卡；第二，操作反馈要连贯，避免状态切换带来的明显卡顿；第三，存档读档过程要稳定，保证关键进度不丢失；第四，界面应具备基本可读性和操作一致性，方便展示与测试。"),
    ("h2", "3.4 课题重难点分析"),
    ("p", "根据任务书内容，课题重难点主要包括动作手感优化、战斗与资源循环设计、敌人行为差异化实现以及探索系统与存档系统的联动。其中，角色控制与状态切换的连贯性直接决定游戏体验，而场景对象状态、能力解锁状态与区域发现状态的同步保存则是保证系统完整性的关键。"),

    ("h1", "4 系统总体设计"),
    ("h2", "4.1 总体架构设计"),
    ("p", "系统采用模块化组织方式，将核心逻辑拆分为玩家模块、敌人模块、成长模块、探索模块、界面模块和存档模块。玩家模块负责输入响应、状态机和战斗逻辑；敌人模块负责巡逻、追击、攻击、受伤和死亡逻辑；成长模块负责等级、经验、装备与属性计算；探索模块负责能力解锁、区域发现与门控交互；界面模块负责状态展示与存档操作；存档模块负责跨场景数据持久化。"),
    ("h2", "4.2 场景设计"),
    ("p", "项目主要包含 MainMenu、MainScene 和 BossScene 三个场景。MainMenu 用于流程入口与开始界面；MainScene 用于承载主体探索流程、装备获取、普通敌人与能力解锁；BossScene 用于展示阶段性挑战与战斗内容。这样的结构能够让游戏原型形成较完整的关卡层级。"),
    ("h2", "4.3 数据流设计"),
    ("p", "玩家输入首先作用于角色状态机，再由状态机控制移动速度、碰撞逻辑、攻击行为和动画表现。战斗结果会改变角色生命值、敌人状态和成长数据；成长数据变化又会影响角色最终属性。界面系统通过读取玩家属性、等级、背包和装备信息完成显示，存档系统则统一收集这些运行数据并在读档时回放到对应模块。"),
    ("h2", "4.4 功能模块划分"),
    ("p", "从代码结构来看，项目脚本主要分布在 Player、Combat、Progression、Save、UI、Audio、Water 和敌人状态脚本等目录中。这样的划分方式与功能边界基本一致，便于在后续扩展时保持模块职责稳定。"),

    ("h1", "5 系统详细设计与实现"),
    ("h2", "5.1 角色控制系统实现"),
    ("p", "玩家脚本 Player 继承基础实体类，并在 Awake 阶段初始化 PlayerStateMachine。系统当前定义了待机、移动、跳跃、空中、冲刺、贴墙滑行、墙跳、近战攻击、远程攻击、受击、回声施法与回声换位等状态。Update 方法负责驱动当前状态更新，并检测冲刺、远程攻击与回声换位输入。"),
    ("p", "在位移方面，角色支持地面移动、跳跃、双段跳、地面冲刺、空中冲刺和墙面交互。能力是否可用由解锁状态控制，例如双段跳和空中冲刺需要在获取对应能力后才可使用。这种设计使成长系统与动作系统直接形成耦合。"),
    ("h2", "5.2 战斗系统实现"),
    ("p", "近战攻击通过攻击检测点和圆形碰撞检测实现。系统利用 Physics2D.OverlapCircleAll 搜索攻击范围内的目标，并通过 IDamageable 接口调用统一受击逻辑。远程攻击则通过生成 WitchThrowing 投射物完成，投射物负责记录方向、伤害值和可命中层级。"),
    ("p", "为了保证动作反馈的连贯性，玩家在攻击状态受到伤害时会先进入伤害缓冲，再在适当时机切换到受击状态。此外，系统还实现了冲刺无敌时间、接触伤害冷却和攻击伤害去重等机制，以减少重复判定问题。"),
    ("h2", "5.3 敌人 AI 与 Boss 实现"),
    ("p", "敌人系统以 EnemyStateMachine 为核心，当前至少包含普通敌人、虫类敌人和迷你 Boss 三类对象。不同敌人都具备待机、移动、攻击、受伤和死亡等基础状态，并通过各自状态脚本表现不同行为差异，例如攻击距离、速度参数和动作节奏。"),
    ("p", "这种状态机结构便于在不破坏主框架的前提下增加新的敌人类型，也符合任务书中对于敌人 AI 与阶段性战斗内容的要求。"),
    ("h2", "5.4 地图探索与能力解锁实现"),
    ("p", "能力类型通过 PlayerAbilityType 枚举统一定义，包括 Dash、DoubleJump、WallJump、AirDash、EchoSwap 和 RangedAttack 六类。AbilityPickup 负责能力拾取，AbilityGate 负责能力门控。玩家具备对应能力后，门控对象会解除阻挡碰撞并切换视觉状态。"),
    ("p", "项目还实现了回声换位和回声通道机制。玩家在解锁 EchoSwap 后能够生成回声实体并与之换位，使空间探索与战斗操作获得更多变化。该机制有效提升了地图设计的层次感。"),
    ("p", "区域发现系统通过 UndiscoveredAreaManager 与 DiscoverableArea 配合实现，使用名为 undiscovered 的 Tilemap 作为未探索区域遮罩。玩家进入新的发现区后，系统移除对应遮罩瓦片，并记录区域编号，读档后再恢复发现结果。"),
    ("h2", "5.5 角色成长与背包装备实现"),
    ("p", "成长系统由 PlayerProgression、PlayerStats、PlayerInventory 和 PlayerEquipment 四个模块构成。PlayerProgression 负责等级、经验和金币；PlayerStats 根据基础属性、等级加成和装备加成计算最终生命值、攻击力、移动速度、跳跃力度和冲刺速度。"),
    ("p", "背包系统负责道具收集与移除，装备系统则按槽位管理穿戴物品。当前项目使用 Charm 和 Relic 两类槽位。装备切换后会通过事件通知刷新属性数据和界面信息，从而形成可见的成长反馈。"),
    ("h2", "5.6 界面系统实现"),
    ("p", "UI_Inventory 脚本使用纯代码动态创建菜单画布、状态页、背包页、装备页和设置页。界面支持打开时暂停游戏、鼠标拖动窗口、阻断场景输入以及在设置页中直接进行三档位存档管理。"),
    ("p", "从界面内容来看，状态页展示等级、经验、金币和生命值，背包页与装备页用于管理道具和属性说明，设置页则提供保存、读取和删除功能。这一设计与参考论文中“模块化功能展示”的写法较为一致。"),
    ("h2", "5.7 存档系统实现"),
    ("p", "SaveSystem 采用单例方式跨场景常驻，默认支持三个档位。系统保存的数据包括场景名称、玩家坐标、当前生命值、等级、经验、金币、已解锁能力、背包内容、已装备道具、已发现区域以及已移除场景对象。"),
    ("p", "在读档过程中，系统先加载目标场景，再在场景加载结束后恢复玩家状态与各模块数据。对于已被拾取的能力道具或已被击败的敌人，系统通过“场景名 + 层级路径 + 组件类型”的唯一标识进行隐藏恢复，以保证场景状态前后一致。"),
    ("h2", "5.8 音频与辅助表现实现"),
    ("p", "BGMManager 负责背景音乐常驻播放、切换和音量控制，ParallaxBackground、水体脚本和雪景特效脚本则用于补充环境表现。虽然这些内容不是系统主干，但它们提升了整体原型的完成度。"),

    ("h1", "6 系统测试"),
    ("h2", "6.1 测试环境"),
    ("p", "项目测试主要在 Windows 环境与 Unity 编辑器下完成，开发版本为 Unity 2022.3.12f1c1。测试方式以人工黑盒测试和模块联调为主，重点验证主流程是否完整、关键数据是否正确同步以及存档读档是否可靠。"),
    ("h2", "6.2 功能测试"),
    ("table_test", ""),
    ("h2", "6.3 测试结果分析"),
    ("p", "测试结果表明，系统已实现较完整的游戏原型闭环。角色控制、战斗、能力解锁、背包装备、区域发现和存档读档能够稳定协同工作。部分内容如地图规模、敌人行为复杂度和视觉反馈仍有继续完善空间，但不影响本课题作为毕业设计成果进行展示。"),

    ("h1", "7 总结与展望"),
    ("h2", "7.1 工作总结"),
    ("p", "本文结合当前 Unity 项目，完成了一个二维类银河恶魔城游戏原型的需求分析、系统设计、关键功能实现与测试整理。项目通过有限状态机完成角色与敌人行为组织，通过 Tilemap 和能力门控构建探索循环，通过成长系统和存档系统形成长期游玩支持，较好体现了游戏开发类毕业设计的综合性。"),
    ("h2", "7.2 后续展望"),
    ("p", "后续可以从以下几个方向继续完善：一是扩展关卡规模和地图复杂度，丰富回访体验；二是增加更多敌人类型与 Boss 阶段机制；三是加入 NPC、剧情和任务系统，强化叙事表达；四是进一步优化音效、动画和特效表现；五是增强数据驱动能力，使新装备、新能力和新敌人能够通过更少脚本修改完成配置。"),

    ("h1", "参考文献"),
    ("refs", ""),
    ("h1", "致谢"),
    ("acks", ""),
]


def set_run_font(run, name, size, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def add_field(run, instruction):
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def build_doc():
    doc = Document()
    doc.settings.odd_and_even_pages_header_footer = True

    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(3.5)
    sec.bottom_margin = Cm(3.0)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.5)
    sec.header_distance = Cm(2.8)
    sec.footer_distance = Cm(2.2)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = Pt(22)
    normal.paragraph_format.first_line_indent = Pt(24)

    h1 = doc.styles["Heading 1"]
    h1.font.name = "黑体"
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)
    h1.paragraph_format.line_spacing = Pt(22)

    h2 = doc.styles["Heading 2"]
    h2.font.name = "黑体"
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(6)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.line_spacing = Pt(22)

    h3 = doc.styles["Heading 3"]
    h3.font.name = "黑体"
    h3._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.paragraph_format.space_before = Pt(6)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.line_spacing = Pt(22)

    return doc


def add_title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(TITLE)
    set_run_font(r, "黑体", 18, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"学    生：{STUDENT}，{COLLEGE}")
    set_run_font(r, "仿宋", 12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"指导教师：{TEACHER}，{COLLEGE}")
    set_run_font(r, "仿宋", 12)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run("【摘要】")
    set_run_font(r, "黑体", 12, True)
    for text in ABSTRACT_CN:
        add_body_paragraph(doc, text)
    add_keyword_line(doc, "【关键词】", KEYWORDS_CN, "黑体", "宋体")
    doc.add_page_break()


def add_english_abstract(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("Design and Implementation of a Unity-based Metroidvania Game")
    set_run_font(r, "Times New Roman", 16, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Student: {STUDENT}, The {COLLEGE}")
    set_run_font(r, "Times New Roman", 12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Teacher: {TEACHER}, The {COLLEGE}")
    set_run_font(r, "Times New Roman", 12)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run("[Abstract]")
    set_run_font(r, "Times New Roman", 12, True)
    for text in ABSTRACT_EN:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.line_spacing = Pt(22)
        r = p.add_run(text)
        set_run_font(r, "Times New Roman", 12)
    add_keyword_line(doc, "[Key words]", KEYWORDS_EN, "Times New Roman", "Times New Roman")
    doc.add_page_break()


def add_keyword_line(doc, label, value, label_font, value_font):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = Pt(22)
    r1 = p.add_run(label)
    set_run_font(r1, label_font, 12, True)
    r2 = p.add_run(value)
    set_run_font(r2, value_font, 12)


def add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(22)
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r, "宋体", 12)


def add_toc(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run("目录")
    set_run_font(r, "黑体", 18, True)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run()
    add_field(run, r'TOC \o "1-3" \h \z \u')
    doc.add_page_break()


def add_body_title(doc):
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(3.5)
    sec.bottom_margin = Cm(3.0)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.5)
    sec.header_distance = Cm(2.8)
    sec.footer_distance = Cm(2.2)
    sec.different_first_page_header_footer = False

    odd = sec.header.paragraphs[0]
    odd.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ro = odd.add_run()
    add_field(ro, 'STYLEREF "Heading 1"')
    set_run_font(ro, "宋体", 9)

    even = sec.even_page_header.paragraphs[0]
    even.alignment = WD_ALIGN_PARAGRAPH.CENTER
    re = even.add_run(TITLE)
    set_run_font(re, "宋体", 9)

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = footer.add_run("第")
    set_run_font(r1, "宋体", 9)
    r2 = footer.add_run()
    add_field(r2, "PAGE")
    set_run_font(r2, "宋体", 9)
    r3 = footer.add_run("页（共")
    set_run_font(r3, "宋体", 9)
    r4 = footer.add_run()
    add_field(r4, "NUMPAGES")
    set_run_font(r4, "宋体", 9)
    r5 = footer.add_run("页）")
    set_run_font(r5, "宋体", 9)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(TITLE)
    set_run_font(r, "黑体", 18, True)


def style_table(table):
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.first_line_indent = Pt(0)
                p.paragraph_format.line_spacing = Pt(18)
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_run_font(run, "宋体", 10.5)


def add_requirement_table(doc):
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "功能模块"
    hdr[1].text = "主要需求"
    hdr[2].text = "项目对应实现"
    rows = [
        ("角色控制", "移动、跳跃、冲刺、墙滑、墙跳、受击反馈", "Player 与 PlayerStateMachine"),
        ("战斗系统", "近战、远程、接触伤害、无敌帧", "攻击检测、投射物与受击逻辑"),
        ("敌人系统", "普通敌人、虫类敌人、迷你 Boss 状态切换", "EnemyStateMachine 及各敌人状态类"),
        ("探索系统", "能力门控、区域发现、特殊通道", "AbilityGate、DiscoverableArea、EchoPassage"),
        ("成长系统", "经验、等级、金币、装备、属性加成", "PlayerProgression、PlayerStats、PlayerEquipment"),
        ("存档系统", "三档位保存、读取、删除、场景状态恢复", "SaveSystem 与 SaveData"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v
    style_table(table)


def add_test_table(doc):
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "测试项"
    hdr[1].text = "测试内容"
    hdr[2].text = "预期结果"
    hdr[3].text = "测试结果"
    rows = [
        ("角色移动", "行走、跳跃、冲刺、墙跳与空中动作", "状态切换正确，动作响应正常", "通过"),
        ("战斗系统", "近战攻击、远程投射与受击反馈", "命中判定正确，伤害逻辑正常", "通过"),
        ("能力解锁", "拾取能力后进入对应门控区域", "门控解锁，能力可正常使用", "通过"),
        ("背包装备", "拾取、装备、卸下与属性刷新", "界面与属性数据同步更新", "通过"),
        ("区域发现", "进入新区域并重新读档", "遮罩清除与区域记录恢复正确", "通过"),
        ("存档读档", "三档位保存、读取、删除", "数据完整保存并可稳定恢复", "通过"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v
    style_table(table)


def add_sections(doc):
    for kind, text in SECTIONS:
        if kind == "h1":
            doc.add_paragraph(text, style="Heading 1")
        elif kind == "h2":
            doc.add_paragraph(text, style="Heading 2")
        elif kind == "h3":
            doc.add_paragraph(text, style="Heading 3")
        elif kind == "p":
            add_body_paragraph(doc, text)
        elif kind == "table_req":
            add_requirement_table(doc)
        elif kind == "table_test":
            add_test_table(doc)
        elif kind == "refs":
            for item in REFERENCES:
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Pt(0)
                p.paragraph_format.line_spacing = Pt(18)
                r = p.add_run(item)
                set_run_font(r, "宋体", 10.5)
        elif kind == "acks":
            for item in ACKS:
                add_body_paragraph(doc, item)


def main():
    doc = build_doc()
    add_title_page(doc)
    add_english_abstract(doc)
    add_toc(doc)
    add_body_title(doc)
    add_sections(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
