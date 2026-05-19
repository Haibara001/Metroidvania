# -*- coding: utf-8 -*-
import json, os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def make_docx(data):
    doc = Document()
    for s in doc.sections:
        s.page_width = Cm(21.0)
        s.page_height = Cm(29.7)
        s.top_margin = Cm(3.5)
        s.bottom_margin = Cm(3.0)
        s.left_margin = Cm(3.0)
        s.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.line_spacing = Pt(22)
    style.paragraph_format.first_line_indent = Pt(24)

    def new_para(text='', bold=False, size=12, font_name='宋体', alignment=None,
                 first_indent=None, space_before=0, space_after=0, line_spacing=22):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = Pt(line_spacing)
        if first_indent is not None:
            pf.first_line_indent = Pt(first_indent)
        elif alignment == WD_ALIGN_PARAGRAPH.CENTER:
            pf.first_line_indent = Pt(0)
        else:
            pf.first_line_indent = Pt(24)
        if alignment:
            p.alignment = alignment
        pf.space_before = Pt(space_before)
        pf.space_after = Pt(space_after)
        if text:
            run = p.add_run(text)
            run.font.size = Pt(size)
            run.font.name = font_name
            run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            run.bold = bold
        return p

    def h1(text):
        new_para(text, bold=True, size=15, font_name='黑体', alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=0, space_before=6, space_after=6)

    def h2(text):
        new_para(text, bold=True, size=14, font_name='黑体', first_indent=0, space_before=6, space_after=6)

    def h3(text):
        new_para(text, bold=True, size=12, font_name='黑体', first_indent=0, space_before=6, space_after=6)

    def body(text):
        new_para(text)

    def body_ni(text):
        new_para(text, first_indent=0)

    def fig(text):
        new_para(text, bold=True, size=10.5, font_name='黑体', alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=0, space_before=6, space_after=6)

    def empty():
        new_para('', first_indent=0)

    def page_break():
        doc.add_page_break()

    def kw_line(label, value):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = Pt(22)
        p.paragraph_format.first_line_indent = Pt(24)
        r1 = p.add_run(label)
        r1.font.size = Pt(12)
        r1.font.name = '宋体'
        r1.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r1.bold = True
        r2 = p.add_run(value)
        r2.font.size = Pt(12)
        r2.font.name = '宋体'
        r2.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    def right(text):
        new_para(text, alignment=WD_ALIGN_PARAGRAPH.RIGHT, first_indent=0)

    # ========== COVER ==========
    for _ in range(6):
        empty()
    new_para(data['university'], bold=True, size=26, font_name='黑体', alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=0)
    empty(); empty()
    new_para('本科毕业论文', bold=True, size=22, font_name='黑体', alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=0)
    empty(); empty()
    new_para(data['title'], bold=True, size=18, font_name='黑体', alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=0)
    for _ in range(4):
        empty()
    for line in [
        '学    院：' + data['college'],
        '专    业：' + data['major'],
        '学    号：' + data['class'],
        '姓    名：' + data['author'],
        '指导老师：' + data['advisor'],
    ]:
        new_para(line, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=0)
    empty()
    new_para(data['date'], size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=0)
    page_break()

    # ========== ABSTRACT CN ==========
    h1('摘  要')
    for line in data['abstract_cn']:
        body(line)
    empty()
    kw_line('关键词：', data['keywords_cn'])
    empty(); empty()

    # ========== ABSTRACT EN ==========
    h1('Abstract')
    for line in data['abstract_en']:
        body(line)
    empty()
    kw_line('Keywords: ', data['keywords_en'])
    page_break()

    # ========== TOC ==========
    h1('目  录')
    body('（请在Word中使用 引用-目录 功能自动生成目录）')
    page_break()

    # ========== SECTIONS ==========
    for sec in data['sections']:
        t = sec.get('type', 'body')
        txt = sec.get('text', '')
        if t == 'h1':
            h1(txt)
        elif t == 'h2':
            h2(txt)
        elif t == 'h3':
            h3(txt)
        elif t == 'body':
            body(txt)
        elif t == 'body_ni':
            body_ni(txt)
        elif t == 'fig':
            fig(txt)
        elif t == 'empty':
            empty()
        elif t == 'page_break':
            page_break()
        elif t == 'kw_line':
            kw_line(sec.get('label', ''), sec.get('value', ''))

    # Save
    out = os.path.join(os.path.dirname(__file__), '毕业论文-基于Unity引擎的类银河恶魔城游戏的设计与实现.docx')
    doc.save(out)
    print(f'Saved: {out}')

if __name__ == '__main__':
    with open(os.path.join(os.path.dirname(__file__), 'thesis_content.json'), 'r', encoding='utf-8') as f:
        data = json.load(f)
    make_docx(data)
